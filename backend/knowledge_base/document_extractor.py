"""Safe text extraction for local mission-reference documents."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_DOCUMENT_BYTES = 8 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 60_000


class DocumentExtractionError(ValueError):
    """Raised when an uploaded reference document cannot be read safely."""


def extract_document_text(filename: str, raw_bytes: bytes) -> str:
    """Return clean, bounded text from a supported reference document."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_DOCUMENT_EXTENSIONS))
        raise DocumentExtractionError(f"Unsupported reference format. Use: {allowed}.")
    if not raw_bytes:
        raise DocumentExtractionError("The uploaded document is empty.")
    if len(raw_bytes) > MAX_DOCUMENT_BYTES:
        raise DocumentExtractionError("Reference documents must be 8 MB or smaller.")
    try:
        if suffix in {".txt", ".md"}:
            text = raw_bytes.decode("utf-8")
        elif suffix == ".pdf":
            from pypdf import PdfReader
            text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(raw_bytes)).pages)
        else:
            from docx import Document
            document = Document(BytesIO(raw_bytes))
            text = "\n".join([p.text for p in document.paragraphs] + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
    except Exception as exc:
        raise DocumentExtractionError("The document could not be read. It may be encrypted or malformed.") from exc
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not cleaned:
        raise DocumentExtractionError("No readable text was found in this document.")
    return cleaned[:MAX_EXTRACTED_CHARACTERS]
